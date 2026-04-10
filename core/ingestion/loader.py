"""
Document loading and basic text extraction.
Ingestion keeps format decoding minimal; semantic parsing belongs to analyzers.
"""
from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
import logging

from .config import LoaderConfig
from .types import DocumentRef, LoadException

log = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """
    Normalized document content after loading.
    """
    # normalized means read the file, extract the text if possible and store the raw_bytes
    raw_bytes: bytes
    extracted_text: Optional[str] = None
    pages_count: Optional[int] = None
    
    @property
    def text_length(self) -> int:
        """Length of extracted text or 0."""
        return len(self.extracted_text) if self.extracted_text else 0


class DocumentLoader:
    """
    Load and extract text from various document formats.
    """
    
    def __init__(self, config: LoaderConfig):
        self.config = config
    
    def load(self, file_ref: DocumentRef) -> LoadedDocument:
        """
        Load document and extract text.
        
        Args:
            file_ref: Reference to document file
        
        Returns:
            LoadedDocument with raw bytes and extracted text
        
        Raises:
            LoadException: If loading fails
        """
        if not file_ref.real_path.exists():
            raise LoadException(f"File not found: {file_ref.real_path}")
        
        try:
            raw_bytes = file_ref.real_path.read_bytes()
        except Exception as e:
            raise LoadException(f"Failed to read file {file_ref.real_path}: {e}")
        
        # Extract text based on format
        extracted_text = None
        pages_count = None
        
        try:
            if file_ref.format == "pdf":
                extracted_text, pages_count = self._load_pdf(raw_bytes, file_ref.real_path)
            
            elif file_ref.format == "html" or file_ref.format == "htm":
                extracted_text = self._load_html(raw_bytes)
            
            elif file_ref.format == "txt":
                extracted_text = self._load_text(raw_bytes)
            
            elif file_ref.format == "docx":
                extracted_text = self._load_docx(raw_bytes)
            
            elif file_ref.format == "md":
                extracted_text = self._load_markdown(raw_bytes)
            
            elif file_ref.format == "xml":
                extracted_text = self._load_xml(raw_bytes)
            
            elif file_ref.format == "json":
                extracted_text = self._load_json(raw_bytes)

            elif file_ref.format == "doc":
                extracted_text = self._load_doc(file_ref.real_path, raw_bytes)

            # Enforce max text length
            if extracted_text and self.config.max_text_length > 0:
                if len(extracted_text) > self.config.max_text_length:
                    log.warning(
                        f"Truncating text for {file_ref.logical_path} "
                        f"({len(extracted_text)} → {self.config.max_text_length} chars)"
                    )
                    extracted_text = extracted_text[:self.config.max_text_length]
        
        except Exception as e:
            log.warning(f"Failed to extract text from {file_ref.logical_path}: {e}")
            # Don't fail the whole ingestion, just skip text extraction
        
        return LoadedDocument(
            raw_bytes=raw_bytes,
            extracted_text=extracted_text,
            pages_count=pages_count,
        )

    def _load_pdf(self, data: bytes, path: Path) -> tuple[Optional[str], Optional[int]]:
        """
        Extract text from PDF.
        
        Handles malformed PDFs gracefully with fallback strategies.
        
        Returns: (text, pages_count)
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            log.warning("pypdf not installed, skipping PDF text extraction")
            return None, None
        
        try:
            from io import BytesIO
            pdf = PdfReader(BytesIO(data))
            pages_count = len(pdf.pages)
            
            if not self.config.pdf_extract_text:
                return None, pages_count
            
            max_pages = self.config.pdf_max_pages or len(pdf.pages)
            max_pages = min(max_pages, len(pdf.pages))
            
            texts = []
            extraction_errors = 0
            
            for i, page in enumerate(pdf.pages[:max_pages]):
                try:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
                except Exception as e:
                    # Log PDF-specific errors at debug level - they're common with malformed PDFs
                    if "startxref" in str(e).lower():
                        log.debug(f"Malformed PDF content at page {i} from {path}: {e}")
                    else:
                        log.debug(f"Failed to extract page {i} from {path}: {e}")
                    extraction_errors += 1
            
            result_text = "\n".join(texts) if texts else None
            
            # Log summary if we had significant extraction errors
            if extraction_errors > max_pages // 2:
                log.debug(f"PDF extraction: {max_pages - extraction_errors}/{max_pages} pages successful from {path}")
            
            return result_text, pages_count
        
        except Exception as e:
            log.debug(f"Failed to read PDF {path}: {e}")
            return None, None
    
    def _load_html(self, data: bytes) -> Optional[str]:
        """Decode HTML bytes without semantic parsing."""
        return self._decode_text(data, "html")
    
    def _load_text(self, data: bytes) -> Optional[str]:
        """Load plain text file."""
        return self._decode_text(data, "text")
    
    def _load_docx(self, data: bytes) -> Optional[str]:
        """Extract text from DOCX."""
        try:
            from docx import Document
            from io import BytesIO
        except ImportError:
            log.warning("python-docx not installed, skipping DOCX text extraction")
            return None
        
        try:
            doc = Document(BytesIO(data))
            texts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            return "\n".join(texts) if texts else None
        
        except Exception as e:
            log.warning(f"Failed to extract text from DOCX: {e}")
            return None
    
    def _load_markdown(self, data: bytes) -> Optional[str]:
        """Decode markdown bytes without semantic parsing."""
        return self._decode_text(data, "markdown")
    
    def _load_xml(self, data: bytes) -> Optional[str]:
        """Decode XML bytes without semantic parsing."""
        return self._decode_text(data, "xml")
    
    def _load_json(self, data: bytes) -> Optional[str]:
        """Decode JSON bytes without semantic parsing or reformatting."""
        return self._decode_text(data, "json")

    def _decode_text(self, data: bytes, format_name: str) -> Optional[str]:
        """Decode bytes using configured fallbacks for text-like formats."""
        try:
            for encoding in self.config.encoding_fallback:
                try:
                    text = data.decode(encoding)
                    return text if text.strip() else None
                except (UnicodeDecodeError, LookupError):
                    continue

            return data.decode("utf-8", errors="replace")
        except Exception as e:
            log.warning(f"Failed to decode {format_name} file: {e}")
            return None

    def _load_doc(self, path: Path, raw_bytes: bytes) -> Optional[str]:
        """Extract text from a legacy .doc file using Word COM automation."""
        try:
            import pythoncom
            import win32com.client
        except ImportError:
            log.warning("pywin32 not installed, skipping DOC text extraction")
            return None

        import tempfile
        import shutil

        tmp_dir = Path(tempfile.mkdtemp(prefix="ingest_doc_"))
        word = None
        document = None
        initialized = False
        text = None
        try:
            pythoncom.CoInitialize()
            initialized = True

            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            tmp_path = tmp_dir / path.name
            tmp_path.write_bytes(raw_bytes)

            document = word.Documents.Open(
                str(tmp_path.resolve()), ReadOnly=True, AddToRecentFiles=False
            )
            raw_text = document.Content.Text
            text = raw_text.strip() if raw_text and raw_text.strip() else None

        except Exception as e:
            log.warning(f"Failed to extract text from DOC {path.name}: {e}")
        finally:
            # Close Word before deleting temp files to release file locks
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            if initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
            shutil.rmtree(tmp_dir, ignore_errors=True)

        return text
